"""All PyMuPDF (fitz) document operations.

Coordinate convention: everything on the wire is in PDF points with the origin
at the TOP-LEFT of the (rotated) page — this matches both pdf.js viewports at
scale 1 and PyMuPDF page coordinates, so no conversion is needed.
"""
import base64
import io
import logging
import math
import statistics
import zipfile
from pathlib import Path

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

import text_engine

# ---------------------------------------------------------------- helpers

BASE14 = {
    ("helv", False, False): "helv",
    ("helv", True, False): "hebo",
    ("helv", False, True): "heit",
    ("helv", True, True): "hebi",
    ("tiro", False, False): "tiro",
    ("tiro", True, False): "tibo",
    ("tiro", False, True): "tiit",
    ("tiro", True, True): "tibi",
    ("cour", False, False): "cour",
    ("cour", True, False): "cobo",
    ("cour", False, True): "coit",
    ("cour", True, True): "cobi",
}


def fontname(family: str, bold: bool = False, italic: bool = False) -> str:
    fam = family if family in ("helv", "tiro", "cour") else "helv"
    return BASE14[(fam, bool(bold), bool(italic))]


def hex2rgb(value: str, default=(0, 0, 0)) -> tuple:
    try:
        v = value.lstrip("#")
        if len(v) == 3:
            v = "".join(c * 2 for c in v)
        return tuple(int(v[i : i + 2], 16) / 255 for i in (0, 2, 4))
    except Exception:
        return default


def parse_ranges(spec: str, page_count: int) -> list[int]:
    """Parse a 1-based range string like '1-3,5,8-' into 0-based page indices."""
    if not spec or spec.strip().lower() == "all":
        return list(range(page_count))
    out: list[int] = []
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, _, b = part.partition("-")
            start = int(a) if a.strip() else 1
            end = int(b) if b.strip() else page_count
        else:
            start = end = int(part)
        for p in range(start, end + 1):
            idx = p - 1
            if 0 <= idx < page_count and idx not in out:
                out.append(idx)
    if not out:
        raise ValueError("Empty page range")
    return out


def parse_range_groups(spec: str, page_count: int) -> list[list[int]]:
    """Each comma-separated segment becomes its own group (used for split)."""
    groups = []
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        groups.append(parse_ranges(part, page_count))
    if not groups:
        raise ValueError("Empty split specification")
    return groups


def _save_over(doc: fitz.Document, path: Path):
    tmp = path.with_suffix(".tmp.pdf")
    doc.save(str(tmp), garbage=3, deflate=True)
    doc.close()
    tmp.replace(path)


def page_count(data: bytes) -> int:
    with fitz.open(stream=data, filetype="pdf") as doc:
        return doc.page_count


# ---------------------------------------------------------------- annotations

_num = text_engine.num  # tolerant float: None/NaN/garbage -> default


def _finite(*vals) -> bool:
    return all(isinstance(v, (int, float)) and math.isfinite(float(v))
               for v in vals)


def _rect(a: dict) -> fitz.Rect:
    if not _finite(a.get("x"), a.get("y"), a.get("w"), a.get("h")):
        raise ValueError("invalid geometry")
    return fitz.Rect(a["x"], a["y"], a["x"] + a["w"], a["y"] + a["h"])


def _insert_textbox_fit(page: fitz.Page, rect: fitz.Rect, text: str, *,
                        font: str, size: float, color: tuple, align: int):
    """insert_textbox that grows the rect / shrinks the font until text fits."""
    rc = page.insert_textbox(rect, text, fontname=font, fontsize=size,
                             color=color, align=align, lineheight=1.25)
    if rc >= 0:
        return
    # grow the box downward by the reported deficit
    grown = fitz.Rect(rect.x0, rect.y0, rect.x1, rect.y1 - rc + size)
    rc = page.insert_textbox(grown, text, fontname=font, fontsize=size,
                             color=color, align=align, lineheight=1.25)
    if rc >= 0:
        return
    fs = size
    while fs > 5:
        fs *= 0.85
        rc = page.insert_textbox(grown, text, fontname=font, fontsize=fs,
                                 color=color, align=align, lineheight=1.25)
        if rc >= 0:
            return


ALIGN = {"left": 0, "center": 1, "right": 2, "justify": 3}


def _apply_text(page: fitz.Page, a: dict):
    rect = _rect(a)
    _insert_textbox_fit(
        page, rect, a.get("text", ""),
        font=fontname(a.get("fontFamily", "helv"), a.get("bold"), a.get("italic")),
        size=max(4.0, _num(a.get("fontSize"), 14)),
        color=hex2rgb(a.get("color") or "#000000"),
        align=ALIGN.get(a.get("align", "left"), 0),
    )
    if a.get("url"):
        page.insert_link({"kind": fitz.LINK_URI, "from": rect, "uri": a["url"]})


def _apply_markup(page: fitz.Page, a: dict):
    rects = [fitz.Rect(r["x"], r["y"], r["x"] + r["w"], r["y"] + r["h"])
             for r in a.get("rects", [])
             if _finite(r.get("x"), r.get("y"), r.get("w"), r.get("h"))]
    rects = [r for r in rects if not r.is_empty]
    if not rects:
        return
    quads = [r.quad for r in rects]
    kind = a["type"]
    if kind == "highlight":
        annot = page.add_highlight_annot(quads=quads)
        annot.set_opacity(_num(a.get("opacity"), 0.45))
    elif kind == "underline":
        annot = page.add_underline_annot(quads=quads)
    else:
        annot = page.add_strikeout_annot(quads=quads)
    annot.set_colors(stroke=hex2rgb(a.get("color") or "#ffd400"))
    annot.update()


def _apply_ink(page: fitz.Page, a: dict):
    pts = [(float(p[0]), float(p[1])) for p in a.get("points", [])
           if isinstance(p, (list, tuple)) and len(p) == 2 and _finite(p[0], p[1])]
    if len(pts) < 2:
        return
    annot = page.add_ink_annot([pts])
    annot.set_colors(stroke=hex2rgb(a.get("color") or "#e11d48"))
    annot.set_border(width=max(0.2, _num(a.get("width"), 2)))
    annot.set_opacity(_num(a.get("opacity"), 1))
    annot.update()


def _apply_shape(page: fitz.Page, a: dict):
    shape = page.new_shape()
    stroke = hex2rgb(a.get("stroke") or "#e11d48")
    fill = hex2rgb(a["fill"]) if a.get("fill") else None
    width = max(0.2, _num(a.get("strokeWidth"), 2))
    opacity = _num(a.get("opacity"), 1)
    kind = a["type"]
    if kind == "rect":
        shape.draw_rect(_rect(a))
    elif kind == "ellipse":
        shape.draw_oval(_rect(a))
    elif kind in ("line", "arrow"):
        if not _finite(a.get("x1"), a.get("y1"), a.get("x2"), a.get("y2")):
            raise ValueError("invalid geometry")
        p1 = fitz.Point(a["x1"], a["y1"])
        p2 = fitz.Point(a["x2"], a["y2"])
        shape.draw_line(p1, p2)
        if kind == "arrow":
            angle = math.atan2(p2.y - p1.y, p2.x - p1.x)
            head = max(9.0, width * 4.5)
            for off in (math.radians(150), -math.radians(150)):
                tip = fitz.Point(
                    p2.x + head * math.cos(angle + off),
                    p2.y + head * math.sin(angle + off),
                )
                shape.draw_line(p2, tip)
    shape.finish(color=stroke, fill=fill, width=width,
                 stroke_opacity=opacity, fill_opacity=opacity,
                 lineCap=1, lineJoin=1)
    shape.commit(overlay=True)


def _apply_image(page: fitz.Page, a: dict):
    src = a.get("src", "")
    if "," in src:
        src = src.split(",", 1)[1]
    src += "=" * ((4 - len(src) % 4) % 4)
    data = base64.b64decode(src)
    rect = _rect(a)
    page.insert_image(rect, stream=data, rotate=int(_num(a.get("rotate"), 0)) % 360,
                      keep_proportion=False, overlay=True)
    if a.get("url"):
        page.insert_link({"kind": fitz.LINK_URI, "from": rect, "uri": a["url"]})


def _apply_note(page: fitz.Page, a: dict):
    if not _finite(a.get("x"), a.get("y")):
        raise ValueError("invalid geometry")
    annot = page.add_text_annot(fitz.Point(a["x"], a["y"]),
                                a.get("text", ""), icon="Comment")
    annot.set_colors(stroke=hex2rgb(a.get("color") or "#ffd400"))
    annot.update()


def _apply_link(page: fitz.Page, a: dict):
    page.insert_link({"kind": fitz.LINK_URI, "from": _rect(a),
                      "uri": a.get("url", "")})


APPLIERS = {
    "text": _apply_text,
    "highlight": _apply_markup,
    "underline": _apply_markup,
    "strikeout": _apply_markup,
    "ink": _apply_ink,
    "rect": _apply_shape,
    "ellipse": _apply_shape,
    "line": _apply_shape,
    "arrow": _apply_shape,
    "image": _apply_image,
    "note": _apply_note,
    "link": _apply_link,
}


def bake_annotations(path: Path, annotations: list[dict]) -> tuple[list[str], list[int]]:
    """Apply the client's pending edit layer permanently into the PDF.

    Returns (layout warnings, changed page numbers). Text rewrites go
    through the paragraph engine — the paragraph is re-derived server-side
    from its id so geometry can never go stale — and may reflow content
    onto the following pages, which is why the changed pages are reported.
    """
    doc = fitz.open(str(path))
    by_page: dict[int, list[dict]] = {}
    for a in annotations:
        pno = int(_num(a.get("page"), -1))
        if 0 <= pno < doc.page_count and a.get("type"):
            by_page.setdefault(pno, []).append(a)

    pool = text_engine.FontPool(doc)
    warnings: list[str] = []
    changed: set[int] = set()
    for pno, items in by_page.items():
        changed.add(pno)
        # 1) free block transforms (move / resize / rotate): islands that
        #    redact in place and redraw elsewhere without shifting the page.
        #    A block op whose geometry is untouched is really a plain rewrite
        #    and is folded into the reflow pass below instead.
        edits = [a for a in items if a.get("type") == "textedit"]
        blocks = [a for a in items if a.get("type") == "textblock"]
        if blocks:
            reflow, free = text_engine.split_block_ops(doc, pno, blocks)
            edits = reflow + edits
            if free:
                w, ch = text_engine.apply_block_ops(doc, pno, free, pool)
                warnings += w
                changed |= ch
        # 2) text rewrites next: their redactions would otherwise wipe
        #    annotations placed over the same area
        if edits:
            w, ch = text_engine.apply_paragraph_edits(doc, pno, edits, pool)
            warnings += w
            changed |= ch
        # 3) everything else in creation order (page fetched after the
        #    edits: a reflow may have appended pages to the document).
        #    One malformed object must not fail the whole batch — skip it
        #    and tell the user, otherwise the client queue can never drain.
        page = doc[pno]
        for a in items:
            fn = APPLIERS.get(a["type"])
            if not fn:
                continue
            try:
                fn(page, a)
            except Exception as exc:
                warnings.append(
                    f"Page {pno + 1}: skipped one {a['type']} object "
                    f"({exc})")

    pool.finalize()
    _save_over(doc, path)
    return warnings, sorted(changed)


def get_paragraphs(path: Path, pno: int) -> list[dict]:
    """Editable paragraphs for one page (click-to-edit overlay)."""
    doc = fitz.open(str(path))
    try:
        if not (0 <= pno < doc.page_count):
            return []
        return text_engine.paragraphs_public(doc, pno)
    finally:
        doc.close()


def add_watermark(path: Path, *, text: str, color: str, opacity: float,
                  font_size: float | None, rotate: float, pages: list[int]):
    doc = fitz.open(str(path))
    rgb = hex2rgb(color, default=(0.7, 0.7, 0.7))
    for pno in pages:
        if not (0 <= pno < doc.page_count):
            continue
        page = doc[pno]
        r = page.rect
        fs = font_size or max(24.0, min(r.width, r.height) * 0.11)
        tl = fitz.get_text_length(text, fontname="hebo", fontsize=fs)
        center = fitz.Point(r.x0 + r.width / 2, r.y0 + r.height / 2)
        origin = fitz.Point(center.x - tl / 2, center.y + fs * 0.35)
        matrix = fitz.Matrix(1, 1).prerotate(rotate)
        page.insert_text(origin, text, fontname="hebo", fontsize=fs,
                         color=rgb, fill_opacity=opacity, overlay=True,
                         morph=(center, matrix))
    _save_over(doc, path)


# ---------------------------------------------------------------- forms

_WIDGET_KINDS = {}          # filled lazily: fitz constants -> wire names


def _widget_kind(ftype: int) -> str | None:
    if not _WIDGET_KINDS:
        _WIDGET_KINDS.update({
            fitz.PDF_WIDGET_TYPE_TEXT: "text",
            fitz.PDF_WIDGET_TYPE_CHECKBOX: "checkbox",
            fitz.PDF_WIDGET_TYPE_COMBOBOX: "choice",
            fitz.PDF_WIDGET_TYPE_LISTBOX: "choice",
        })
    return _WIDGET_KINDS.get(ftype)


def list_form_fields(path: Path) -> list[dict]:
    """All fillable AcroForm widgets (text, checkbox, combo/list boxes).

    Coordinates are page points, top-left origin, like everything else.
    """
    # ponytail: radio groups and signature widgets are skipped — add radio
    # via widget.on_state() if documents with radio forms show up
    doc = fitz.open(str(path))
    try:
        out = []
        for pno in range(doc.page_count):
            for w in doc[pno].widgets() or []:
                kind = _widget_kind(w.field_type)
                if kind is None:
                    continue
                r = w.rect
                value = w.field_value
                if kind == "checkbox":
                    value = value not in (None, "", "Off", False)
                out.append({
                    "xref": w.xref,
                    "name": w.field_name or f"field-{w.xref}",
                    "page": pno,
                    "x": r.x0, "y": r.y0, "w": r.width, "h": r.height,
                    "type": kind,
                    "value": value if value is not None else "",
                    "options": (w.choice_values or []) if kind == "choice" else [],
                    "fontSize": w.text_fontsize or 0,
                })
        return out
    finally:
        doc.close()


def set_form_fields(path: Path, values: dict[int, object]) -> int:
    """Set widget values by xref. Returns how many fields were updated."""
    doc = fitz.open(str(path))
    n = 0
    for pno in range(doc.page_count):
        for w in doc[pno].widgets() or []:
            if w.xref not in values or _widget_kind(w.field_type) is None:
                continue
            v = values[w.xref]
            if w.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                w.field_value = bool(v)
            else:
                w.field_value = "" if v is None else str(v)
            w.update()
            n += 1
    if n:
        _save_over(doc, path)
    else:
        doc.close()
    return n


# ---------------------------------------------------------------- reading

def get_outline(path: Path) -> list[dict]:
    """Existing bookmarks, or a heading heuristic when the PDF has none."""
    doc = fitz.open(str(path))
    try:
        toc = doc.get_toc(simple=True) or []
        if toc:
            return [{"level": lvl, "title": title, "page": max(0, pno - 1)}
                    for lvl, title, pno in toc]
        # Heuristic: lines whose font size clearly exceeds the body size.
        sizes: list[float] = []
        lines: list[tuple[int, float, str]] = []
        for pno in range(min(doc.page_count, 80)):
            d = doc[pno].get_text("dict")
            for block in d.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    txt = "".join(s["text"] for s in spans).strip()
                    if not txt:
                        continue
                    size = max(s["size"] for s in spans)
                    sizes.append(size)
                    lines.append((pno, size, txt))
        if not sizes:
            return []
        body = statistics.median(sizes)
        heads = [(p, s, t) for p, s, t in lines
                 if s >= body * 1.3 and 2 < len(t) < 90 and not t.endswith(".")]
        if not heads:
            return []
        uniq_sizes = sorted({round(s, 1) for _, s, _ in heads}, reverse=True)
        tier = {s: min(i + 1, 3) for i, s in enumerate(uniq_sizes)}
        return [{"level": tier[round(s, 1)], "title": t, "page": p}
                for p, s, t in heads][:200]
    finally:
        doc.close()


def list_notes(path: Path) -> list[dict]:
    """All baked sticky-note annotations in the document."""
    doc = fitz.open(str(path))
    try:
        out = []
        for pno in range(doc.page_count):
            for annot in doc[pno].annots(types=[fitz.PDF_ANNOT_TEXT]) or []:
                info = annot.info
                out.append({
                    "xref": annot.xref, "page": pno,
                    "x": annot.rect.x0, "y": annot.rect.y0,
                    "text": info.get("content", ""),
                    "modified": info.get("modDate", ""),
                })
        return out
    finally:
        doc.close()


def _find_note(doc: fitz.Document, xref: int):
    for pno in range(doc.page_count):
        # keep the page object: the annot only weak-references its parent page,
        # so returning a fresh doc[pno] would leave the annot orphaned
        page = doc[pno]
        for annot in page.annots(types=[fitz.PDF_ANNOT_TEXT]) or []:
            if annot.xref == xref:
                return page, annot
    return None, None


def update_note(path: Path, xref: int, text: str) -> bool:
    doc = fitz.open(str(path))
    page, annot = _find_note(doc, xref)
    if annot is None:
        doc.close()
        return False
    annot.set_info(content=text)
    annot.update()
    _save_over(doc, path)
    return True


def delete_note(path: Path, xref: int) -> bool:
    doc = fitz.open(str(path))
    page, annot = _find_note(doc, xref)
    if annot is None:
        doc.close()
        return False
    page.delete_annot(annot)
    _save_over(doc, path)
    return True


# ---------------------------------------------------------------- page ops

def rotate_pages(path: Path, pages: list[int], degrees: int):
    doc = fitz.open(str(path))
    for pno in pages:
        if 0 <= pno < doc.page_count:
            page = doc[pno]
            page.set_rotation((page.rotation + degrees) % 360)
    _save_over(doc, path)


def delete_pages(path: Path, pages: list[int]) -> int:
    doc = fitz.open(str(path))
    keep = [p for p in pages if 0 <= p < doc.page_count]
    if len(keep) >= doc.page_count:
        doc.close()
        raise ValueError("Cannot delete every page of the document")
    doc.delete_pages(keep)
    n = doc.page_count
    _save_over(doc, path)
    return n


def insert_blank_page(path: Path, after: int) -> int:
    """Insert one blank page after 0-based index `after` (-1 = at the front).
    The new page copies the size of the page it follows (or page 1)."""
    doc = fitz.open(str(path))
    ref = doc[max(0, min(doc.page_count - 1, after))]
    doc.new_page(pno=after + 1, width=ref.rect.width, height=ref.rect.height)
    n = doc.page_count
    _save_over(doc, path)
    return n


def duplicate_pages(path: Path, pages: list[int]) -> int:
    doc = fitz.open(str(path))
    # process descending so earlier indices stay valid
    for pno in sorted(set(pages), reverse=True):
        if 0 <= pno < doc.page_count:
            doc.fullcopy_page(pno, pno + 1)
    n = doc.page_count
    _save_over(doc, path)
    return n


def reorder_pages(path: Path, order: list[int]):
    doc = fitz.open(str(path))
    if sorted(order) != list(range(doc.page_count)):
        doc.close()
        raise ValueError("Order must be a permutation of all pages")
    doc.select(order)
    _save_over(doc, path)


def extract_pages(path: Path, pages: list[int]) -> bytes:
    doc = fitz.open(str(path))
    valid = [p for p in pages if 0 <= p < doc.page_count]
    if not valid:
        doc.close()
        raise ValueError("No valid pages selected")
    doc.select(valid)
    buf = doc.tobytes(garbage=3, deflate=True)
    doc.close()
    return buf


def split_document(path: Path, groups: list[list[int]]) -> list[bytes]:
    out = []
    for group in groups:
        out.append(extract_pages(path, group))
    return out


def merge_documents(paths: list[Path]) -> bytes:
    merged = fitz.open()
    for p in paths:
        with fitz.open(str(p)) as src:
            merged.insert_pdf(src)
    buf = merged.tobytes(garbage=3, deflate=True)
    merged.close()
    return buf


# ---------------------------------------------------------------- export

def export_images(path: Path, pages: list[int], fmt: str, dpi: int, password: str = None) -> list[tuple[str, bytes]]:
    """Render pages to PNG/JPG. Returns [(filename, bytes)]."""
    doc = fitz.open(str(path))
    if doc.is_encrypted and password:
        doc.authenticate(password)
    ext = "png" if fmt == "png" else "jpg"
    out = []
    for pno in pages:
        if not (0 <= pno < doc.page_count):
            continue
        pix = doc[pno].get_pixmap(dpi=dpi, alpha=False)
        data = pix.tobytes(output="png" if fmt == "png" else "jpeg",
                           jpg_quality=88)
        out.append((f"page-{pno + 1}.{ext}", data))
    doc.close()
    return out


def zip_bytes(entries: list[tuple[str, bytes]]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in entries:
            zf.writestr(name, data)
    return buf.getvalue()


def compress(data: bytes, quality: int = 70) -> bytes:
    """Optimize PDF by aggressively downsampling and re-compressing embedded raster images."""
    import logging
    from PIL import Image
    logger = logging.getLogger(__name__)
    
    # Map quality to max image dimensions to keep file sizes small while maintaining quality
    if quality <= 50: # Extreme
        max_dim = 1000
    elif quality <= 70: # Recommended
        max_dim = 1500
    else: # Less
        max_dim = 2000

    with fitz.open(stream=data, filetype="pdf") as doc:
        # Iterate through pages and optimize images
        processed_xrefs = set()
        for page in doc:
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                if xref in processed_xrefs:
                    continue
                processed_xrefs.add(xref)
                
                try:
                    base_image = doc.extract_image(xref)
                    if not base_image:
                        continue
                    img_bytes = base_image["image"]
                    
                    # Don't compress tiny images (e.g. less than 15 KB) to save CPU
                    if len(img_bytes) < 15360:
                        continue
                        
                    im = Image.open(io.BytesIO(img_bytes))
                    
                    # Convert transparent images to white background to prevent text visibility loss
                    if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
                        bg = Image.new("RGB", im.size, (255, 255, 255))
                        if im.mode == "RGBA":
                            bg.paste(im, mask=im.split()[3])
                        else:
                            bg.paste(im.convert("RGBA"), mask=im.convert("RGBA").split()[3])
                        im = bg
                    elif im.mode == "CMYK":
                        im = im.convert("RGB")
                    elif im.mode != "RGB":
                        im = im.convert("RGB")
                        
                    # Downsample if image exceeds max dimension
                    if max(im.size) > max_dim:
                        im.thumbnail((max_dim, max_dim), Image.Resampling.LANCZOS)
                        
                    # Save as optimized JPEG
                    out_img = io.BytesIO()
                    im.save(out_img, format="JPEG", quality=quality, optimize=True)
                    compressed_bytes = out_img.getvalue()
                    
                    # Only swap the image in the PDF if it actually reduces size
                    if len(compressed_bytes) < len(img_bytes):
                        page.replace_image(xref, stream=compressed_bytes)
                except Exception as exc:
                    logger.warning(f"Failed to compress image xref {xref}: {exc}")
                    
        # Apply standard PDF garbage collection, stream deflating, and cleaning
        out = io.BytesIO()
        doc.save(out, garbage=4, deflate=True, clean=True)
        compressed_data = out.getvalue()
        
        if len(compressed_data) >= len(data):
            return data
        return compressed_data


def compress_to_target_size(data: bytes, target_size: int) -> bytes:
    """Try different quality levels to get PDF size below or equal to target_size."""
    best_data = None
    best_size = float('inf')
    
    # Try compressing at different quality steps from high to low
    for q in [90, 75, 60, 45, 30, 20]:
        try:
            compressed = compress(data, quality=q)
            c_size = len(compressed)
            
            # If we met the target size, we can stop and return this
            if c_size <= target_size:
                return compressed
                
            # Otherwise keep track of the smallest size we achieved
            if c_size < best_size:
                best_size = c_size
                best_data = compressed
        except Exception:
            continue
            
    # Return the best (smallest) data we achieved if none got below the target,
    # ensuring it's actually smaller than the original.
    if best_data is not None and len(best_data) < len(data):
        return best_data
    return data


def summarize_pdf(pdf_path, mode: str = "medium") -> str:
    """
    Summarize a PDF using Google Gemini AI.
    Falls back to a basic local extraction if no GEMINI_API_KEY is set.
    Supports modes: 'short', 'medium', and 'detailed'.
    """
    import os
    import fitz

    mode = mode.lower()
    if mode not in ("short", "medium", "detailed"):
        mode = "medium"

    # Extract text from PDF
    doc = fitz.open(str(pdf_path))
    pages_text = []
    for page_idx, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            pages_text.append(f"--- Page {page_idx + 1} ---\n{text}")
    doc.close()

    full_text = "\n\n".join(pages_text)
    if not full_text.strip():
        return "# Summary\n\nThe document does not contain sufficient readable text for summarization."

    # Truncate to ~30k chars to fit Gemini context limits comfortably
    if len(full_text) > 30000:
        full_text = full_text[:30000] + "\n\n[... Document truncated for processing ...]"

    openrouter_error = None
    groq_error = None
    openai_error = None
    gemini_error = None

    # Try OpenRouter API first
    openrouter_key = os.environ.get("OPENROUTER_API_KEY", "")
    if openrouter_key and openrouter_key != "your-openrouter-api-key-here":
        try:
            summary = _summarize_with_openrouter(full_text, mode, openrouter_key)
            return summary
        except Exception as exc:
            openrouter_error = str(exc)
            logger.warning(f"OpenRouter summarization failed: {exc}, trying Groq/OpenAI/Gemini/local fallbacks...")

    # Try Groq API next
    groq_key = os.environ.get("GROQ_API_KEY", "")
    if groq_key and groq_key != "your-groq-api-key-here":
        try:
            summary = _summarize_with_groq(full_text, mode, groq_key)
            return summary
        except Exception as exc:
            groq_error = str(exc)
            logger.warning(f"Groq summarization failed: {exc}, trying OpenAI/Gemini/local fallbacks...")

    # Try OpenAI API next
    openai_key = os.environ.get("OPENAI_API_KEY", "")
    if openai_key and openai_key != "your-openai-api-key-here":
        try:
            summary = _summarize_with_openai(full_text, mode, openai_key)
            return summary
        except Exception as exc:
            openai_error = str(exc)
            logger.warning(f"OpenAI summarization failed: {exc}, trying Gemini/local fallbacks...")

    # Try Gemini API next
    api_key = os.environ.get("GEMINI_API_KEY", "")
    if api_key and api_key != "your-gemini-api-key-here":
        try:
            summary = _summarize_with_gemini(full_text, mode, api_key)
            # If it fell back to local inside the helper, it will mention "Local analysis" or "Fallback"
            if "Local analysis" not in summary:
                return summary
        except Exception as exc:
            gemini_error = str(exc)
            logger.warning(f"Gemini API call failed: {exc}")

    # Fallback to local
    local_summary = _summarize_local(full_text, mode)

    # Append any API errors to help the user debug their keys
    error_notes = []
    if openrouter_key and openrouter_key != "your-openrouter-api-key-here":
        error_notes.append(f"- **OpenRouter API error:** `{openrouter_error or 'Unknown error or empty response'}`")
    if groq_key and groq_key != "your-groq-api-key-here":
        error_notes.append(f"- **Groq API error:** `{groq_error or 'Unknown error or empty response'}`")
    if openai_key and openai_key != "your-openai-api-key-here":
        error_notes.append(f"- **OpenAI API error:** `{openai_error or 'Unknown error or empty response'}`")
    if api_key and api_key != "your-gemini-api-key-here":
        error_notes.append(f"- **Gemini API error:** `{gemini_error or 'Unknown error or empty response'}`")

    if error_notes:
        err_msg = "\n\n---\n### ⚠️ AI API Key Issues Detected\nAI APIs failed to generate a summary. Falling back to the local text-processing engine. Please check your keys or usage limits:\n" + "\n".join(error_notes)
        return local_summary + err_msg

    return local_summary


def _summarize_with_openrouter(text: str, mode: str, api_key: str) -> str:
    """Use OpenRouter API (via OpenAI SDK compatibility) to generate an intelligent summary."""
    import openai

    mode_instructions = {
        "short": """Generate a SHORT executive summary (2-3 paragraphs max).
Include:
- A brief overview paragraph
- 3-5 key bullet points
- Any critical dates, figures, or names mentioned
Keep it concise and to the point.""",

        "medium": """Generate a MEDIUM-length summary.
Include:
- An executive overview (2-3 paragraphs)
- 6-10 key insights as bullet points
- Section-by-section breakdown if the document has clear sections
- Important dates, figures, monetary values, and percentages
- Key names, organizations, and entities mentioned
- Any action items or next steps found""",

        "detailed": """Generate a DETAILED comprehensive summary.
Include:
- A thorough executive overview (3-5 paragraphs)
- 12-20 key insights as bullet points covering all major topics
- Complete section-by-section summaries preserving the document structure
- ALL important dates, deadlines, and timelines
- ALL monetary figures, percentages, and key metrics
- ALL names, organizations, titles, and contact information
- ALL action items, obligations, requirements, and next steps
- Any legal terms, conditions, or compliance requirements
- Technical specifications or data points if present"""
    }

    prompt = f"""You are a professional document analyst. Analyze the following PDF document text and produce a structured summary in clean Markdown format.

{mode_instructions[mode]}

Format your response as clean Markdown with:
- Use # for the main title "Document Summary"
- Use ## for major sections (Overview, Key Insights, Section Summaries, Important Dates, Key Figures, Names & Organizations, Action Items)
- Use ### for subsections
- Use bullet points (- ) for lists
- Use checkbox format (- [ ] ) for action items
- Use **bold** for emphasis on critical information
- Use *italics* for document metadata

IMPORTANT: Only include information that is actually present in the document. Do not fabricate or assume any information.

Here is the document text:

{text}"""

    client = openai.OpenAI(api_key=api_key, base_url="https://openrouter.ai/api/v1")
    response = client.chat.completions.create(
        model="google/gemini-2.5-flash",
        messages=[
            {"role": "user", "content": prompt}
        ],
        max_tokens=3000,
        temperature=0.3
    )
    return response.choices[0].message.content.strip()


def _summarize_with_groq(text: str, mode: str, api_key: str) -> str:
    """Use Groq API (via OpenAI SDK compatibility) to generate an intelligent summary."""
    import openai

    mode_instructions = {
        "short": """Generate a SHORT executive summary (2-3 paragraphs max).
Include:
- A brief overview paragraph
- 3-5 key bullet points
- Any critical dates, figures, or names mentioned
Keep it concise and to the point.""",

        "medium": """Generate a MEDIUM-length summary.
Include:
- An executive overview (2-3 paragraphs)
- 6-10 key insights as bullet points
- Section-by-section breakdown if the document has clear sections
- Important dates, figures, monetary values, and percentages
- Key names, organizations, and entities mentioned
- Any action items or next steps found""",

        "detailed": """Generate a DETAILED comprehensive summary.
Include:
- A thorough executive overview (3-5 paragraphs)
- 12-20 key insights as bullet points covering all major topics
- Complete section-by-section summaries preserving the document structure
- ALL important dates, deadlines, and timelines
- ALL monetary figures, percentages, and key metrics
- ALL names, organizations, titles, and contact information
- ALL action items, obligations, requirements, and next steps
- Any legal terms, conditions, or compliance requirements
- Technical specifications or data points if present"""
    }

    prompt = f"""You are a professional document analyst. Analyze the following PDF document text and produce a structured summary in clean Markdown format.

{mode_instructions[mode]}

Format your response as clean Markdown with:
- Use # for the main title "Document Summary"
- Use ## for major sections (Overview, Key Insights, Section Summaries, Important Dates, Key Figures, Names & Organizations, Action Items)
- Use ### for subsections
- Use bullet points (- ) for lists
- Use checkbox format (- [ ] ) for action items
- Use **bold** for emphasis on critical information
- Use *italics* for document metadata

IMPORTANT: Only include information that is actually present in the document. Do not fabricate or assume any information.

Here is the document text:

{text}"""

    client = openai.OpenAI(api_key=api_key, base_url="https://api.groq.com/openai/v1")
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()


def _summarize_with_openai(text: str, mode: str, api_key: str) -> str:
    """Use OpenAI API to generate an intelligent summary."""
    import openai

    mode_instructions = {
        "short": """Generate a SHORT executive summary (2-3 paragraphs max).
Include:
- A brief overview paragraph
- 3-5 key bullet points
- Any critical dates, figures, or names mentioned
Keep it concise and to the point.""",

        "medium": """Generate a MEDIUM-length summary.
Include:
- An executive overview (2-3 paragraphs)
- 6-10 key insights as bullet points
- Section-by-section breakdown if the document has clear sections
- Important dates, figures, monetary values, and percentages
- Key names, organizations, and entities mentioned
- Any action items or next steps found""",

        "detailed": """Generate a DETAILED comprehensive summary.
Include:
- A thorough executive overview (3-5 paragraphs)
- 12-20 key insights as bullet points covering all major topics
- Complete section-by-section summaries preserving the document structure
- ALL important dates, deadlines, and timelines
- ALL monetary figures, percentages, and key metrics
- ALL names, organizations, titles, and contact information
- ALL action items, obligations, requirements, and next steps
- Any legal terms, conditions, or compliance requirements
- Technical specifications or data points if present"""
    }

    prompt = f"""You are a professional document analyst. Analyze the following PDF document text and produce a structured summary in clean Markdown format.

{mode_instructions[mode]}

Format your response as clean Markdown with:
- Use # for the main title "Document Summary"
- Use ## for major sections (Overview, Key Insights, Section Summaries, Important Dates, Key Figures, Names & Organizations, Action Items)
- Use ### for subsections
- Use bullet points (- ) for lists
- Use checkbox format (- [ ] ) for action items
- Use **bold** for emphasis on critical information
- Use *italics* for document metadata

IMPORTANT: Only include information that is actually present in the document. Do not fabricate or assume any information.

Here is the document text:

{text}"""

    client = openai.OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()


def _summarize_with_gemini(text: str, mode: str, api_key: str) -> str:
    """Use Google Gemini AI to generate an intelligent summary."""

    mode_instructions = {
        "short": """Generate a SHORT executive summary (2-3 paragraphs max).
Include:
- A brief overview paragraph
- 3-5 key bullet points
- Any critical dates, figures, or names mentioned
Keep it concise and to the point.""",

        "medium": """Generate a MEDIUM-length summary.
Include:
- An executive overview (2-3 paragraphs)
- 6-10 key insights as bullet points
- Section-by-section breakdown if the document has clear sections
- Important dates, figures, monetary values, and percentages
- Key names, organizations, and entities mentioned
- Any action items or next steps found""",

        "detailed": """Generate a DETAILED comprehensive summary.
Include:
- A thorough executive overview (3-5 paragraphs)
- 12-20 key insights as bullet points covering all major topics
- Complete section-by-section summaries preserving the document structure
- ALL important dates, deadlines, and timelines
- ALL monetary figures, percentages, and key metrics
- ALL names, organizations, titles, and contact information
- ALL action items, obligations, requirements, and next steps
- Any legal terms, conditions, or compliance requirements
- Technical specifications or data points if present"""
    }

    prompt = f"""You are a professional document analyst. Analyze the following PDF document text and produce a structured summary in clean Markdown format.

{mode_instructions[mode]}

Format your response as clean Markdown with:
- Use # for the main title "Document Summary"
- Use ## for major sections (Overview, Key Insights, Section Summaries, Important Dates, Key Figures, Names & Organizations, Action Items)
- Use ### for subsections
- Use bullet points (- ) for lists
- Use checkbox format (- [ ] ) for action items
- Use **bold** for emphasis on critical information
- Use *italics* for document metadata

IMPORTANT: Only include information that is actually present in the document. Do not fabricate or assume any information.

Here is the document text:

{text}"""

    # Try the google-generativeai SDK first
    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)

        # Handle safety-blocked or empty responses
        if response.candidates:
            candidate = response.candidates[0]
            if candidate.content and candidate.content.parts:
                result = candidate.content.parts[0].text
                if result and result.strip():
                    return result.strip()
            # Check if blocked by safety
            if hasattr(candidate, 'finish_reason'):
                finish = str(candidate.finish_reason)
                if "SAFETY" in finish.upper():
                    logger.warning(f"Gemini response blocked by safety filter: {finish}")
                    return _summarize_local(text, mode)

        # Fallback: try .text property directly
        if hasattr(response, 'text') and response.text:
            return response.text.strip()

        logger.warning("Gemini returned empty response, falling back to local")
        return _summarize_local(text, mode)

    except ImportError:
        logger.warning("google-generativeai not installed, trying REST API")
    except Exception as exc:
        logger.warning(f"Gemini SDK call failed: {exc}")

    # Fallback: direct REST API call
    try:
        import urllib.request
        import json as json_mod

        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
        payload = json_mod.dumps({
            "contents": [{"parts": [{"text": prompt}]}]
        }).encode("utf-8")

        req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json_mod.loads(resp.read().decode("utf-8"))

        candidates = data.get("candidates", [])
        if candidates:
            parts = candidates[0].get("content", {}).get("parts", [])
            if parts and parts[0].get("text", "").strip():
                return parts[0]["text"].strip()

        logger.warning("Gemini REST API returned empty response, falling back to local")
        return _summarize_local(text, mode)

    except Exception as exc:
        logger.warning(f"Gemini REST API call also failed: {exc}, falling back to local summarization")
        return _summarize_local(text, mode)


def _summarize_local(full_text: str, mode: str) -> str:
    """Fallback local NLP summarizer using word-frequency sentence ranking."""
    import re
    import collections

    # Clean text helper
    def clean_sentences(text_content):
        raw_sents = re.split(r'(?<=[.!?])\s+', text_content)
        cleaned = []
        for s in raw_sents:
            s_clean = s.strip().replace("\n", " ")
            s_clean = re.sub(r'\s+', ' ', s_clean)
            if len(s_clean) > 20 and not s_clean.startswith("[") and not s_clean.startswith("---"):
                cleaned.append(s_clean)
        return cleaned

    all_sentences = clean_sentences(full_text)
    if not all_sentences:
        return "# Summary\n\nThe document does not contain sufficient readable text for summarization."

    # Word frequency scoring
    stopwords = {
        "the", "a", "an", "and", "or", "but", "if", "then", "else", "when", "at", "by", "for",
        "with", "about", "against", "between", "into", "through", "during", "before", "after",
        "above", "below", "to", "from", "up", "down", "in", "out", "on", "off", "over", "under",
        "again", "further", "once", "here", "there", "all", "any", "both", "each", "few", "more",
        "most", "other", "some", "such", "no", "nor", "not", "only", "own", "same", "so", "than",
        "too", "very", "can", "will", "just", "don", "should", "now", "is", "was", "were",
        "be", "been", "being", "have", "has", "had", "having", "do", "does", "did", "doing",
        "page", "this", "that", "these", "those", "which", "who", "whom", "what", "where",
    }

    words = re.findall(r'\b[a-z]{3,}\b', full_text.lower())
    freq = collections.Counter(w for w in words if w not in stopwords)

    def score_sentence(sent):
        s_words = re.findall(r'\b[a-z]{3,}\b', sent.lower())
        if not s_words:
            return 0
        return sum(freq[w] for w in s_words) / len(s_words)

    scored = [(s, score_sentence(s)) for s in all_sentences]
    scored.sort(key=lambda x: x[1], reverse=True)

    # Mode-based counts
    counts = {
        "short":    (3, 5),
        "medium":   (5, 10),
        "detailed": (8, 18),
    }
    n_exec, n_bullets = counts.get(mode, counts["medium"])

    md = ["# Document Summary", f"*Mode: {mode.capitalize()} · Local analysis (no AI API key configured)*\n"]

    md.append("## Overview")
    top = scored[:n_exec]
    top.sort(key=lambda x: all_sentences.index(x[0]))
    md.append(" ".join(s[0] for s in top))
    md.append("")

    md.append("## Key Insights")
    for s, _ in scored[n_exec:n_exec + n_bullets]:
        md.append(f"- {s}")
    md.append("")

    # Extract dates, figures
    dates = sorted(set(re.findall(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b', full_text, re.I)))
    dates += sorted(set(re.findall(r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b', full_text)))
    figures = sorted(set(re.findall(r'\$[\d,]+(?:\.\d+)?(?:\s*(?:million|billion|trillion))?', full_text, re.I)))
    figures += sorted(set(re.findall(r'\b\d+(?:\.\d+)?%', full_text)))

    if dates:
        md.append("## Important Dates")
        for d in dates[:10]:
            md.append(f"- {d}")
        md.append("")
    if figures:
        md.append("## Key Figures")
        for f in figures[:10]:
            md.append(f"- {f}")
        md.append("")

    return "\n".join(md)


def markdown_to_pdf(md_text: str, output_path: str, source_filename: str = "document.pdf"):
    """Generate a stylized PDF from Markdown text using ReportLab, including running header/footer watermarks."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    import re
    import datetime

    # Clean the source filename
    clean_source = source_filename
    if clean_source.endswith("_summary.pdf"):
        clean_source = clean_source[:-12] + ".pdf"
    elif clean_source.endswith("_summary"):
        clean_source = clean_source[:-8] + ".pdf"
    elif not clean_source.lower().endswith(".pdf"):
        clean_source = clean_source + ".pdf"

    def draw_header_footer(canvas, doc):
        canvas.saveState()
        # Header text
        canvas.setFont('Helvetica-Bold', 7.5)
        canvas.setFillColor(colors.HexColor('#9ca3af')) # gray-400
        canvas.drawString(54, 745, "QT PDF STUDIO — SUMMARY REPORT")
        
        # Header date
        date_str = datetime.date.today().strftime("%m/%d/%Y")
        canvas.setFont('Helvetica', 7.5)
        canvas.drawRightString(doc.pagesize[0] - 54, 745, date_str)
        
        # Header line
        canvas.setStrokeColor(colors.HexColor('#e5e7eb')) # gray-200
        canvas.setLineWidth(0.5)
        canvas.line(54, 735, doc.pagesize[0] - 54, 735)
        
        # Footer line
        canvas.line(54, 60, doc.pagesize[0] - 54, 60)
        
        # Footer source name
        canvas.setFont('Helvetica', 7.5)
        canvas.setFillColor(colors.HexColor('#4b5563')) # gray-600
        canvas.drawString(54, 46, f"Source: {clean_source}")
        
        # Footer page number
        canvas.setFont('Helvetica-Bold', 7.5)
        canvas.drawRightString(doc.pagesize[0] - 54, 46, "PAGE 1 OF 1")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=54, leftMargin=54,
        topMargin=72, bottomMargin=72
    )
    styles = getSampleStyleSheet()
    
    # Custom styles
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#4f46e5'), # Brand Indigo
        spaceAfter=15,
        spaceBefore=10
    )
    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading2'],
        fontSize=13,
        leading=17,
        textColor=colors.HexColor('#1f2937'), # Gray-800
        spaceAfter=8,
        spaceBefore=12
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#374151'), # Gray-700
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=20,
        firstLineIndent=-10,
        spaceAfter=4
    )

    story = []

    lines = md_text.split('\n')
    for line in lines:
        line_str = line.strip()
        if not line_str:
            story.append(Spacer(1, 8))
            continue
            
        # Clean formatting (bold: **text** -> <b>text</b>, italic: *text* -> <i>text</i>)
        line_str = re.sub(r'\*\*(.*?)\*\*|__(.*?)__', r'<b>\1\2</b>', line_str)
        line_str = re.sub(r'\*(.*?)\*|_(.*?)_', r'<i>\1\2</i>', line_str)

        # Parse tags
        if line.startswith('# '):
            text = line[2:].strip()
            story.append(Paragraph(text, h1_style))
        elif line.startswith('## '):
            text = line[3:].strip()
            story.append(Paragraph(text, h2_style))
        elif line.startswith('### '):
            text = line[4:].strip()
            mini_style = ParagraphStyle('Mini', parent=h2_style, fontSize=11, leading=14, textColor=colors.HexColor('#4f46e5'), spaceBefore=8, spaceAfter=4)
            story.append(Paragraph(text, mini_style))
        elif line.startswith('- [ ] ') or line.startswith('- [x] '):
            checked = line.startswith('- [x] ')
            box = "[X]" if checked else "[ ]"
            text = f"<font color='#4f46e5'><b>{box}</b></font> " + line[6:].strip()
            story.append(Paragraph(text, bullet_style))
        elif line.startswith('- ') or line.startswith('* '):
            text = "<font color='#4f46e5'><b>•</b></font> " + line[2:].strip()
            story.append(Paragraph(text, bullet_style))
        else:
            story.append(Paragraph(line_str, body_style))

    doc.build(story, onFirstPage=draw_header_footer, onLaterPages=draw_header_footer)


def markdown_to_docx(md_text: str, output_path: str, source_filename: str = "document.pdf"):
    """Generate a stylized Word Document (DOCX) from Markdown text using python-docx, including headers/footers."""
    import docx
    import re
    from docx.shared import Pt, RGBColor
    import datetime

    # Clean the source filename
    clean_source = source_filename
    if clean_source.endswith("_summary.docx"):
        clean_source = clean_source[:-13] + ".pdf"
    elif clean_source.endswith("_summary"):
        clean_source = clean_source[:-8] + ".pdf"
    elif not clean_source.lower().endswith(".pdf"):
        clean_source = clean_source + ".pdf"

    doc = docx.Document()
    
    # Configure default normal style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x37, 0x41, 0x51) # Gray-700
    
    # Configure Header and Footer in Section
    section = doc.sections[0]
    
    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "" # Reset default paragraph
    hp.paragraph_format.space_after = Pt(8)
    date_str = datetime.date.today().strftime("%m/%d/%Y")
    hrun = hp.add_run(f"QT PDF STUDIO — SUMMARY REPORT\t\t{date_str}")
    hrun.font.name = 'Arial'
    hrun.font.size = Pt(7.5)
    hrun.font.bold = True
    hrun.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf) # gray-400
    
    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "" # Reset default paragraph
    fp.paragraph_format.space_before = Pt(8)
    frun = fp.add_run(f"Source: {clean_source}\t\tPAGE 1 OF 1")
    frun.font.name = 'Arial'
    frun.font.size = Pt(7.5)
    frun.font.bold = True
    frun.font.color.rgb = RGBColor(0x4b, 0x55, 0x63) # gray-600
    
    lines = md_text.split('\n')
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        # Parse headings
        if line.startswith('# '):
            heading_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(heading_text)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5) # Brand Indigo
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(heading_text)
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37) # Gray-800
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(heading_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5) # Brand Indigo
        else:
            is_bullet = False
            is_checkbox = False
            checked = False
            
            if line.startswith('- [ ] ') or line.startswith('- [x] '):
                is_checkbox = True
                checked = line.startswith('- [x] ')
                line_str = line[6:].strip()
            elif line.startswith('- ') or line.startswith('* '):
                is_bullet = True
                line_str = line[2:].strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            
            if is_checkbox:
                box_run = p.add_run("[X]  " if checked else "[ ]  ")
                box_run.bold = True
                box_run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
            elif is_bullet:
                bullet_run = p.add_run("•  ")
                bullet_run.bold = True
                bullet_run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
                
            # Parse bold/italic
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', line_str)
            for part in parts:
                if not part:
                    continue
                if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
                    text_content = part[2:-2]
                    run = p.add_run(text_content)
                    run.bold = True
                elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
                    text_content = part[1:-1]
                    run = p.add_run(text_content)
                    run.italic = True
                else:
                    p.add_run(part)
                    
    doc.save(output_path)



